import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from ..userdatafiles import load_user_data

def generate_cover_letter(output_path: str, body_text: str):
    try:
        CONFIG = load_user_data()
    except Exception as e:
        raise e

    # User details
    name = CONFIG['name']
    address = CONFIG['address']
    email = CONFIG['email']

    # Create a new Word document
    doc = Document()
    
    # Add user name as title (Title)
    title_paragraph = doc.add_paragraph(name, 'Title')
    for run in title_paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0) # Defaults to Blue

    # Add address and date in the same paragraph
    address_date_paragraph = doc.add_paragraph()
    address_date_paragraph.add_run(address).bold = False
    address_date_paragraph.add_run("\nDate: " + datetime.today().strftime("%B %d, %Y")).bold = False

    # Add body text
    doc.add_paragraph("Dear Recruiting Manager,")
    body_split = body_text.split("\n")
    for body in body_split:
        doc.add_paragraph(body)

    # Add closing text
    closing = ("Thank you for your time and consideration. "
               "I look forward to hearing from you soon.")
    doc.add_paragraph(closing)
    
    # Add Sincerely paragraph
    sincerely_paragraph = doc.add_paragraph()
    sincerely_paragraph.add_run("Sincerely,").bold = False
    sincerely_paragraph.add_run("\n"+str(name)).bold = False
    sincerely_paragraph.add_run("\n"+str(email)).bold = False

    # Save the document
    doc.save(output_path + '/coverletter.docx')
